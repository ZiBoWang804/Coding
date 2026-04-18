from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[3]
PROJECT_DIR = ROOT
SUBMISSION_DIR = Path(__file__).resolve().parents[1]
REPORT_PATH = Path(__file__).resolve().parent / "网页主题聚类专题项目实验报告.docx"
BACKUP_PATH = Path(__file__).resolve().parent / "网页主题聚类专题项目实验报告.bak.docx"
FALLBACK_REPORT_PATH = Path(__file__).resolve().parent / "网页主题聚类专题项目实验报告_重建版.docx"
SCREENSHOT_DIR = SUBMISSION_DIR / "screenshots"
DASHBOARD_PATH = PROJECT_DIR / "generated" / "dashboard.json"


def load_dashboard() -> dict:
    return json.loads(DASHBOARD_PATH.read_text(encoding="utf-8"))


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(paragraph, text: str, size: float = 12, bold: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    if "BodyTextCn" not in styles:
        body_style = styles.add_style("BodyTextCn", WD_STYLE_TYPE.PARAGRAPH)
    else:
        body_style = styles["BodyTextCn"]
    body_style.base_style = normal
    body_style.font.name = "宋体"
    body_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body_style.font.size = Pt(12)
    body_style.paragraph_format.line_spacing = 1.5
    body_style.paragraph_format.first_line_indent = Cm(0.74)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="BodyTextCn")
    set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_paragraph_text(paragraph, text, size={1: 16, 2: 14, 3: 12}[level], bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_text(paragraph, text, size=10.5)


def add_picture_block(document: Document, image_path: Path, caption: str, note: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"截图缺失: {image_path}")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image_path), width=Cm(15.2))
    add_caption(document, caption)
    add_body(document, note)


def build_environment_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(11)

    headers = table.rows[0].cells
    headers[0].text = "项目"
    headers[1].text = "说明"
    for cell in headers:
        for paragraph in cell.paragraphs:
            set_paragraph_text(paragraph, paragraph.text, size=11, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        ("操作系统", "Windows 环境，本地使用 Python 3.11.5 运行与调试。"),
        ("后端服务", "Python 标准库 http.server 提供页面与数据接口，入口脚本为 run_app.py。"),
        ("核心依赖", "beautifulsoup4、requests、numpy、scikit-learn，用于采集、特征提取和聚类实验。"),
        ("前端实现", "原生 HTML、CSS、JavaScript，包含首页、主题页、文章页和来源页。"),
        ("聚类方法", "TF-IDF 特征表示结合 KMeans 聚类，对候选 K 值自动比较并选择最佳结果。"),
        ("情感分析", "轻量词典式情感分析，将文章划分为正面、中性、负面，并生成时间线趋势图。"),
        ("启动方式", "在项目根目录执行 python run_app.py，浏览器访问 http://127.0.0.1:8765。"),
        ("部署形态", "支持本地演示，也可导出为静态站点并部署到 Cloudflare Pages。"),
    ]

    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        for idx, cell in enumerate(row):
            for paragraph in cell.paragraphs:
                set_paragraph_text(paragraph, paragraph.text, size=11, bold=False)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_report() -> None:
    dashboard = load_dashboard()
    metrics = dashboard["metrics"]
    tracked = dashboard["tracked_topic"]
    clusters = dashboard["clusters"]
    cluster_labels = "、".join(cluster["label"] for cluster in clusters)
    generated_at = datetime.fromisoformat(dashboard["generated_at"]).strftime("%Y年%m月%d日 %H:%M")

    if REPORT_PATH.exists():
        shutil.copy2(REPORT_PATH, BACKUP_PATH)

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_paragraph_text(title, "网页主题聚类专题项目实验报告", size=18, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    set_paragraph_text(subtitle, f"重构时间：{generated_at}", size=11)

    add_heading(document, "一、题目背景", 1)
    add_body(
        document,
        "本专题项目围绕“网页主题聚类、主题跟踪与情感分析”展开，目标是把采集到的网页文章按照内容相似度自动归并成多个主题簇，并将聚类结果以可视化方式展示给用户。用户不需要逐篇阅读全部文章，只需先浏览聚类结果，再从中选择自己关心的一个主题簇作为后续跟踪对象，这样就能够把数据挖掘课程中的文本聚类方法与一个可直接交互的网页系统结合起来。"
    )
    add_body(
        document,
        f"本次实验使用项目当前生成的数据快照作为结果依据，共纳入 {metrics['article_count']} 篇网页文章、覆盖 {metrics['source_count']} 个站点来源，系统自动比较候选 K 值 {metrics['candidate_ks']} 并选择最佳聚类数 K={metrics['best_k']}。当前聚类结果形成 {metrics['cluster_count']} 个主题方向，涵盖 {cluster_labels} 等内容类别，既能满足课程作业对“聚类结果展示”的要求，也为后续的主题跟踪和情感分析提供了可解释的主题入口。"
    )
    add_body(
        document,
        f"在交互流程上，系统默认优先展示最佳 K 值下的聚类结果，并将每个簇总结为中文主题名称和代表关键词。用户选中某个簇后，系统会自动把该簇关键词作为主题线索继续筛选相关文章，并对其进行情感倾向分析。当前默认关注主题为“{tracked['label']}”，概括性关键词为“{tracked['topic_keyword']}”，跟踪到 {metrics['tracked_article_count']} 篇相关文章，整体情绪判断为“{tracked['dominant_sentiment_label']}”，这正好对应了老师提出的“先聚类，再关注，再看情感”的专题思路。"
    )

    add_heading(document, "二、开发环境", 1)
    add_body(
        document,
        "项目采用前后端分离但不依赖复杂框架的实现方式，便于在课程环境下快速部署、演示和复现实验结果。后端负责数据组织、聚类计算和接口输出，前端负责中文界面呈现、网络关系图、主题卡片、文章详情与情感时间线等交互展示。整体环境配置见表 1。"
    )
    build_environment_table(document)
    add_caption(document, "表1 开发环境与主要技术组成")

    add_heading(document, "三、核心算法流程", 1)
    add_body(
        document,
        "系统算法流程可以概括为“数据采集与样本构建—文本预处理—TF-IDF 向量化—KMeans 聚类—主题命名—主题跟踪—情感分析—前端可视化输出”。该流程既保留了数据挖掘课程中常见的文本聚类主线，也结合网页应用的交互需求，对聚类后主题的解释与展示做了进一步整理。"
    )

    add_heading(document, "3.1 数据采集与样本构建", 2)
    add_body(
        document,
        "项目同时支持演示样本和在线 RSS 抓取两种数据来源。本次报告采用项目内置的演示数据快照，以保证课堂展示时数据稳定、页面内容完整。采集后的样本会统一整理为文章标题、摘要、来源名称、来源网址、发布日期等字段，形成后续聚类和情感分析所需的基础语料库。由于网页来源覆盖体育、人工智能、金融、能源、文旅等多类新闻站点，因此聚类结果能够更明显地体现出不同主题之间的内容差异。"
    )

    add_heading(document, "3.2 文本预处理与特征表示", 2)
    add_body(
        document,
        "在聚类之前，系统会将文章标题与摘要拼接成文本样本，并进行基础清洗，去掉空白噪声和无意义片段。随后使用 TF-IDF 方法把文本转换为向量表示。TF-IDF 的优势在于能够突出某篇文章中更有区分度的词项，同时降低常见词对聚类结果的干扰，因此适合用于新闻网页这一类主题性较强的文本数据。"
    )

    add_heading(document, "3.3 KMeans 聚类与最佳 K 值选择", 2)
    add_body(
        document,
        f"完成向量化之后，系统会对候选聚类数 {metrics['candidate_ks']} 逐一计算聚类结果，并结合轮廓系数与聚类质量分数选择最佳 K 值。当前快照中自动得到的最佳结果为 K={metrics['best_k']}，对应轮廓系数为 {metrics['silhouette_score']:.3f}，聚类质量分数为 {metrics['clustering_quality']:.3f}。这种做法避免了把主题类别完全写死到代码中，而是先通过算法划分文本集合，再在聚类完成后总结每个簇的主题名称和核心关键词。"
    )

    add_heading(document, "3.4 主题命名与用户关注逻辑", 2)
    add_body(
        document,
        "聚类结束后，系统会从每个簇的高频关键词和代表文章中提取最有概括性的主题线索，再生成用户可理解的中文主题名称。例如，一个簇中如果集中出现“光伏、风电、电网”等词，就会被命名为“能源电力”；如果某个簇主要出现“利率、汇率、股价、市场预期”等词，则会被归纳为金融类主题。这样一来，用户看到的不再是抽象的簇编号，而是具有明确语义的主题卡片与网络节点。"
    )
    add_body(
        document,
        "在交互层面，用户可以直接点击首页中的主题簇或关系网中心节点进入主题页。系统会把该主题簇标记为“关注主题”，并将其关键词集合作为后续相关文章检索和情感判断的依据。这个过程实现了从无监督聚类结果到用户个性化关注主题的自然过渡。"
    )

    add_heading(document, "3.5 主题跟踪与情感分析", 2)
    add_body(
        document,
        "当用户确定关注主题后，系统会依据主题关键词与簇中心相似度继续筛选相关文章，并输出与该主题最相关的一批文章。随后，项目采用轻量词典式情感分析方法，统计文章中正向、负向和中性表达的命中情况，给出每篇文章的情绪标签，同时在时间维度上汇总为情感时间线。当前默认关注的“能源电力”主题共筛选出 8 篇文章，其中正面 2 篇、中性 2 篇、负面 4 篇，因此整体情绪判断为偏负面。"
    )

    add_heading(document, "3.6 前端可视化输出", 2)
    add_body(
        document,
        "为了让聚类结果更适合课堂展示，前端使用中文界面将聚类簇网络、来源覆盖、主题卡片、主题详情、文章详情和情感折线图组织成一套完整的浏览路径。关系网视图通过簇中心与文章节点之间的连线展示聚类结构；主题页把代表文章、来源站点和情感走势整合到同一页面；文章页和来源页又分别承担“查看原文内容”和“查看站点来源”的作用，形成完整的从聚类到详情的可视化闭环。"
    )

    document.add_page_break()

    add_heading(document, "四、结果截图展示说明", 1)
    add_body(
        document,
        "为了更直观地说明系统运行效果，下面结合实际页面截图展示项目的主要功能模块。截图全部来自当前项目的真实运行页面，并且按首页总览、关系网、主题卡片、主题详情、文章详情、来源详情、情感折线图的顺序进行说明。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "home-top.png",
        "图1 首页顶部总览与统计指标",
        f"图 1 展示了系统首页的顶部区域，可以看到文章总量、站点来源数量、最佳聚类数等核心指标已经汇总显示出来。这里使用的是当前数据快照的真实统计结果，即 {metrics['article_count']} 篇文章、{metrics['source_count']} 个来源、最佳 K 值为 {metrics['best_k']}。这一部分对应老师要求中的“将聚类结果显示给用户”，用户进入系统后可以先从总体规模和聚类数量上把握当前样本集合。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "network-graph.png",
        "图2 首页关系网可视化",
        "图 2 补充了原报告中缺失的关系网可视化部分。页面中心展示的是主题簇节点，外围分布的是与各簇相关的文章节点，节点之间通过连线反映文章与主题簇的归属关系。用户可以通过鼠标悬停和点击来观察簇的中文名称及文章分布情况，从而直观看到“哪些网页内容被聚在一起、不同簇之间如何区分”，这也是本项目最能体现聚类结果的数据可视化模块。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "cluster-cards.png",
        "图3 聚类主题卡片展示",
        "图 3 展示了首页中的主题卡片区域。系统把每一个聚类簇总结为一个可读的中文主题，并在卡片中给出关键词、代表文章、来源数量和情感倾向，用户可以直接从这些卡片里选择一个主题进行关注。相比单纯输出聚类编号，这种展示方式更符合课程作业对“用户可选择感兴趣类别”的要求，也能提高聚类结果的可解释性。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "topic-detail-crop.png",
        "图4 主题详情页",
        f"图 4 是用户进入主题页后的界面效果。当前页面默认展示关注主题“{tracked['label']}”，同时给出主题关键词“{tracked['topic_keyword']}”、相关文章列表、来源站点和情感分布信息。这个页面承担了“把一个聚类簇转化为用户可持续跟踪的主题”的作用，是从无监督聚类结果过渡到个性化跟踪的关键页面。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "article-detail.png",
        "图5 文章详情页",
        "图 5 展示文章详情页。用户在主题页点击任意文章后，可以进一步查看文章标题、摘要、所属主题、情感标签及原始来源链接。这样既保留了聚类与情感分析的结果，又支持用户回到具体网页内容层面进行核验，避免系统只给出抽象分类而没有具体文章支撑。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "source-detail.png",
        "图6 来源详情页",
        "图 6 为来源详情页，用于集中展示某个站点在当前样本库中的文章覆盖情况与站点链接。用户点击来源模块后，可以继续跳转查看对应网站，方便从“主题—文章—来源站点”三个层次理解数据来源。该页面也说明本项目并非只展示单一数据源，而是聚合了多个新闻网站和资讯平台的网页样本。"
    )

    add_picture_block(
        document,
        SCREENSHOT_DIR / "sentiment-timeline.png",
        "图7 主题情感时间线折线图",
        "图 7 补充了原报告中缺失的情感折线图展示。系统把关注主题下相关文章的正面、中性和负面数量按日期汇总，绘制成时间线折线图，用户可以从图中观察情绪变化趋势。相比静态的情感占比展示，折线图更适合说明“主题情感是如何随时间变化的”，因此更能体现老师要求中的“跟踪主题并了解文章情感动态”。"
    )

    add_body(
        document,
        "综合上述截图可以看出，本项目已经完整实现了从网页采集、文本聚类、主题解释、用户关注到情感分析与结果展示的全流程。尤其是关系网可视化和情感折线图两部分，直接对应了聚类结果展示与主题动态跟踪两个核心要求，也使整套网页系统更适合作为课程专题作业进行演示和答辩。"
    )

    try:
        document.save(REPORT_PATH)
        print(f"报告已生成：{REPORT_PATH}")
    except PermissionError:
        document.save(FALLBACK_REPORT_PATH)
        print(f"目标文件被占用，已另存为：{FALLBACK_REPORT_PATH}")


if __name__ == "__main__":
    create_report()
