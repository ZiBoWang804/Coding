from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_ROOT / "output"
DOC_DIR = OUTPUT_DIR / "doc"
SCREEN_DIR = OUTPUT_DIR / "submission_screens"
REPORT_PATH = DOC_DIR / "网页主题聚类专题项目实验报告.docx"
SUBMISSION_ROOT = PROJECT_ROOT / "submission"
PACKAGE_NAME = "网页主题聚类专题项目"
PACKAGE_DIR = SUBMISSION_ROOT / PACKAGE_NAME
ZIP_BASE = SUBMISSION_ROOT / PACKAGE_NAME


def load_state() -> dict:
    path = PROJECT_ROOT / "generated" / "dashboard.json"
    return json.loads(path.read_text(encoding="utf-8"))


def unique_cluster_labels(state: dict) -> list[str]:
    labels: list[str] = []
    for cluster in state.get("clusters", []):
        label = cluster.get("label", "").strip()
        if label and label not in labels:
            labels.append(label)
    return labels


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        heading = document.styles[name]
        heading.font.name = "微软雅黑"
        heading.font.bold = True
        heading.font.size = Pt(size)
        heading._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)


def add_center_paragraph(document: Document, text: str, size: int, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run(text)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.add_run(text)


def add_figure(document: Document, image_path: Path, caption: str, width_cm: float = 15.6) -> None:
    if not image_path.exists():
        return
    document.add_picture(str(image_path), width=Cm(width_cm))
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.bold = True
    caption_run.font.name = "微软雅黑"
    caption_run.font.size = Pt(10.5)
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build_report(state: dict) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    labels = unique_cluster_labels(state)
    tracked = state["tracked_topic"]
    metrics = state["metrics"]

    document = Document()
    configure_document(document)

    add_center_paragraph(document, "网页主题聚类专题项目实验报告", 20, bold=True)
    add_center_paragraph(document, "数据挖掘专题作业配套文档", 14)
    add_center_paragraph(document, f"生成日期：{state.get('generated_at', '')[:10]}", 11)
    document.add_paragraph()

    document.add_heading("一、实验目的", level=1)
    add_body_paragraph(
        document,
        "本实验围绕“网页主题聚类、用户关注主题跟踪与情感分析”展开，目标是构建一个能够对网页文章进行自动聚类、"
        "向用户展示聚类结果、支持选择感兴趣主题并持续观察该主题后续文章情感变化的可视化系统。项目既满足课程作业的"
        "数据挖掘要求，也尽量保留真实网页分析与交互展示的完整链路。",
    )

    document.add_heading("二、实验环境", level=1)
    env_table = document.add_table(rows=5, cols=2)
    env_table.style = "Table Grid"
    env_rows = [
        ("操作系统", "Windows（本地课程实验环境）"),
        ("主要语言", "Python、HTML、CSS、JavaScript"),
        ("核心算法", "TF-IDF 文本表示 + 本地 numpy KMeans 聚类 + 词典式情感分析"),
        ("运行方式", "在项目目录执行 python run_app.py，浏览器访问 http://127.0.0.1:8765"),
        ("当前快照规模", f"{metrics['article_count']} 篇文章，{metrics['source_count']} 个来源，最佳聚类数 best_k={metrics['best_k']}"),
    ]
    for row, (left, right) in zip(env_table.rows, env_rows):
        row.cells[0].text = left
        row.cells[1].text = right

    document.add_heading("三、项目需求与总体设计", level=1)
    add_body_paragraph(
        document,
        "系统的核心业务流程是：先收集网页或 RSS 文章样本，对文本进行预处理与向量化，然后执行聚类，将内容相似的网页"
        "归入同一主题簇；再将每个簇转化为用户容易理解的中文主题名、关键词和代表文章；最后允许用户把某个簇设为关注"
        "主题，继续查看与该主题相关的新文章及其正面、中性、负面情感变化。",
    )
    add_body_paragraph(
        document,
        f"当前演示快照中，系统自动在候选 K 值 {metrics['candidate_ks']} 中选择 best_k={metrics['best_k']}，形成了"
        f"{metrics['cluster_count']} 个聚类簇，主要涵盖 {('、'.join(labels[:6]))} 等主题方向。当前默认关注主题为“{tracked['label']}”，"
        f"概括性关键词为“{tracked['topic_keyword']}”，已跟踪相关文章 {len(tracked['articles'])} 篇。",
    )

    document.add_heading("四、关键算法与实现说明", level=1)
    document.add_heading("4.1 数据采集与样例库", level=2)
    add_body_paragraph(
        document,
        "项目同时保留了 demo、live、auto 三种数据模式。课堂演示时默认使用内置样例库，以保证网络不稳定时仍能完整展示系统。"
        "当前样例库已经扩充到体育、AI 技术、金融市场、能源电力、医疗科技、文化旅游等多个方向，来源网站数量达到 32 个。",
    )
    document.add_heading("4.2 文本表示与真实聚类", level=2)
    add_body_paragraph(
        document,
        "后端首先对文章标题、摘要和来源提示信息进行统一清洗，然后使用 TF-IDF 提取文本特征。聚类阶段没有再使用固定专题词典"
        "直接归类，而是采用本地 numpy 实现的 KMeans，对候选 K 值进行自动搜索，并结合轮廓系数与整体聚类质量分数选择最优结果。"
        f"当前快照的轮廓系数为 {metrics['silhouette_score']}，聚类质量分数为 {metrics['clustering_quality']}。",
    )
    document.add_heading("4.3 主题命名与关键词提取", level=2)
    add_body_paragraph(
        document,
        "聚类完成后，系统会根据簇内高权重词、代表文章和来源分布提炼中文主题名称，同时给出一个最有概括性的主题关键词。"
        "这样既保留了无监督聚类的真实性，也避免仅向用户展示抽象的簇编号。",
    )
    document.add_heading("4.4 主题跟踪与情感分析", level=2)
    add_body_paragraph(
        document,
        "当用户选择某个簇作为关注主题后，系统会把该簇的关键词与中心特征作为后续跟踪线索，对新文章进行主题相关度筛选。"
        "情感分析部分采用轻量级词典法，将文章划分为正面、中性、负面，并在页面中以情感概览卡片与折线时间线的形式展示变化。",
    )

    document.add_heading("五、界面设计与功能展示", level=1)
    add_bullet(document, "首页展示总览指标、聚类关系图、来源覆盖、聚类卡片、关注主题模块与情感趋势。")
    add_bullet(document, "聚类关系图中的主题簇、文章节点、来源节点均支持点击跳转。")
    add_bullet(document, "主题详情页显示该主题下的网页样本、后续跟踪文章以及关键词和情感结构。")
    add_bullet(document, "文章详情页显示摘要、主题线索、来源与相关文章，并附带来源网站跳转入口。")
    add_bullet(document, "来源详情页汇总某个网站收录的文章与主题覆盖情况，便于解释数据来源。")

    document.add_heading("六、系统运行结果与截图说明", level=1)
    add_body_paragraph(
        document,
        "以下截图来自当前项目运行版本，展示了系统从聚类总览到主题详情、文章详情和来源详情的完整交互链路。",
    )
    add_figure(document, SCREEN_DIR / "home-top.png", "图 1 系统首页总览与核心指标")
    add_body_paragraph(
        document,
        "图 1 展示了系统首页顶部区域。页面左侧是“网页主题聚类专题观察台”的主视觉与系统说明，中央统计卡片展示了文章总量、"
        "来源站点数、聚类数和当前跟踪文章数，右侧控制面板可切换数据模式并重新执行聚类。",
    )
    add_figure(document, SCREEN_DIR / "home-full.png", "图 2 首页完整运行结果")
    add_body_paragraph(
        document,
        "图 2 展示了首页完整滚动页面，其中包含聚类关系图、聚类卡片、当前关注主题的关键词、来源样本、情感结构以及情感时间线，"
        "完整体现了老师要求的“聚类展示 + 主题关注 + 情感观察”三段式功能。",
    )
    add_figure(document, SCREEN_DIR / "topic-detail.png", "图 3 主题详情页")
    add_body_paragraph(
        document,
        f"图 3 以当前默认关注主题“{tracked['label']}”为例，展示该主题的关键词、来源和簇内文章样本，同时继续列出与该主题相关的后续文章，"
        "便于用户进一步跟踪某一专题的发展情况。",
    )
    add_figure(document, SCREEN_DIR / "article-detail.png", "图 4 文章详情页")
    add_body_paragraph(
        document,
        "图 4 展示单篇文章的摘要、关键词、所属主题和相关性分值。用户可以从这里回溯文章与主题之间的联系，也可以继续查看同来源的其他文章。",
    )
    add_figure(document, SCREEN_DIR / "source-detail.png", "图 5 来源详情页")
    add_body_paragraph(
        document,
        "图 5 展示来源网站详情页，页面汇总了该来源的文章数量、主题覆盖和文章列表，并保留来源网站跳转按钮，便于说明网页数据来自哪些站点。",
    )

    document.add_heading("七、测试与验证", level=1)
    test_table = document.add_table(rows=5, cols=3)
    test_table.style = "Table Grid"
    header = test_table.rows[0].cells
    header[0].text = "验证项"
    header[1].text = "结果"
    header[2].text = "说明"
    rows = [
        ("python -m unittest discover -s tests -v", "通过", "单元测试全部通过，当前基线为 4 项测试"),
        ("python -m src.pipeline", "通过", "能够重新生成 generated/dashboard.json"),
        ("node --check web/app.js", "通过", "前端脚本语法检查通过"),
        ("本地服务启动", "通过", "执行 python run_app.py 后可访问本地页面"),
        ("页面截图核验", "通过", "已生成首页、主题页、文章页、来源页截图"),
    ]
    for row, values in zip(test_table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    document.add_heading("八、实验总结与改进方向", level=1)
    add_body_paragraph(
        document,
        "本项目完成了从网页文本聚类到用户关注主题跟踪，再到情感分析与可视化展示的完整闭环。相较于简单的固定分类方案，"
        "当前版本恢复为真实的 TF-IDF + KMeans 聚类，并支持自动选择最佳聚类数，从而更贴合数据挖掘课程实验的要求。",
    )
    add_body_paragraph(
        document,
        "后续仍可继续改进的方向包括：进一步优化中文分词和主题命名质量、引入更高精度的中文情感模型、扩展在线抓取规模与数据更新机制、"
        "以及继续打磨图谱布局算法，使不同主题簇的边界更加稳定和直观。",
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录：项目运行说明", level=1)
    add_bullet(document, r"项目目录：D:\wzb\数据挖掘\web classification")
    add_bullet(document, "启动命令：python run_app.py")
    add_bullet(document, "访问地址：http://127.0.0.1:8765")
    add_bullet(document, "推荐演示模式：demo（内置样例）")
    add_bullet(document, "项目当前输出快照：generated/dashboard.json")

    document.save(REPORT_PATH)


def write_submission_readme(state: dict) -> None:
    metrics = state["metrics"]
    tracked = state["tracked_topic"]
    content = f"""网页主题聚类专题项目提交说明

1. 项目简介
本项目用于完成数据挖掘课程中的“网页聚类 + 用户关注主题跟踪 + 情感分析”专题作业。

2. 当前快照
- 文章数量：{metrics['article_count']}
- 来源数量：{metrics['source_count']}
- 聚类数量：{metrics['cluster_count']}
- 最佳聚类数：{metrics['best_k']}
- 当前默认关注主题：{tracked['label']}（关键词：{tracked['topic_keyword']}）

3. 运行方法
- 进入 project 目录后执行：python run_app.py
- 浏览器访问：http://127.0.0.1:8765

4. 目录说明
- report/：实验项目报告 docx 文件
- screenshots/：运行截图
- project/：项目源码与运行文件
"""
    (PACKAGE_DIR / "README_提交说明.txt").write_text(content, encoding="utf-8")


def safe_rmtree(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)


def copy_submission_files() -> None:
    SUBMISSION_ROOT.mkdir(parents=True, exist_ok=True)
    safe_rmtree(PACKAGE_DIR)
    PACKAGE_DIR.mkdir(parents=True, exist_ok=True)

    report_dir = PACKAGE_DIR / "report"
    screens_dir = PACKAGE_DIR / "screenshots"
    project_dir = PACKAGE_DIR / "project"
    report_dir.mkdir()
    screens_dir.mkdir()
    project_dir.mkdir()

    shutil.copy2(REPORT_PATH, report_dir / REPORT_PATH.name)

    for image in [
        "home-top.png",
        "home-full.png",
        "topic-detail.png",
        "article-detail.png",
        "source-detail.png",
    ]:
        source = SCREEN_DIR / image
        if source.exists():
            shutil.copy2(source, screens_dir / image)

    files_to_copy = [
        "README.md",
        "run_app.py",
        "docs/build_submission_bundle.py",
        "docs/implementation_plan.md",
        "docs/run_capture_submission.cmd",
        "docs/system_flow.md",
        "generated/dashboard.json",
    ]
    dirs_to_copy = ["src", "web", "tests"]

    for relative in files_to_copy:
        source = PROJECT_ROOT / relative
        target = project_dir / relative
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)

    for relative in dirs_to_copy:
        source = PROJECT_ROOT / relative
        target = project_dir / relative
        shutil.copytree(
            source,
            target,
            dirs_exist_ok=True,
            ignore=shutil.ignore_patterns("__pycache__", "*.pyc"),
        )


def build_zip() -> Path:
    zip_path = shutil.make_archive(str(ZIP_BASE), "zip", root_dir=SUBMISSION_ROOT, base_dir=PACKAGE_NAME)
    return Path(zip_path)


def main() -> None:
    state = load_state()
    build_report(state)
    copy_submission_files()
    write_submission_readme(state)
    zip_path = build_zip()
    print(f"Report written to: {REPORT_PATH}")
    print(f"Package directory: {PACKAGE_DIR}")
    print(f"Zip archive: {zip_path}")


if __name__ == "__main__":
    main()
